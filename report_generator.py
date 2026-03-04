import logging
import base64
import os
from io import BytesIO
from typing import Dict, Any, Optional
from datetime import datetime

# docx imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# reportlab imports for PDF
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.colors import HexColor
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Generate Word and PDF reports from KPI analysis with RTL support."""
    
    def __init__(self):
        self.colors_rgb = {
            "primary": (55, 65, 81),         # Charcoal gray
            "accent": (79, 70, 229),         # Modern indigo
            "success": (16, 185, 129),       # Fresh green
            "text_primary": (17, 24, 39),    # Near black
            "text_secondary": (107, 114, 128), # Muted gray
        }
        
        self.docx_colors = {
            key: RGBColor(*rgb) for key, rgb in self.colors_rgb.items()
        }
        
        # Setup PDF fonts if available
        self._setup_pdf_fonts()

    def _setup_pdf_fonts(self):
        """Try to load a font that supports Arabic (like Arial or Tahoma) for PDF."""
        self.arabic_font_name = "Helvetica" # Fallback
        
        # Check standard Windows font paths for arial.ttf or tahoma.ttf
        font_paths = [
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/tahoma.ttf",
            "C:/Windows/Fonts/times.ttf"
        ]
        
        for path in font_paths:
            if os.path.exists(path):
                try:
                    font_name = os.path.basename(path).split(".")[0].capitalize()
                    pdfmetrics.registerFont(TTFont(font_name, path))
                    self.arabic_font_name = font_name
                    logger.info(f"Registered {font_name} for PDF generation.")
                    break
                except Exception as e:
                    logger.warning(f"Failed to load font {path}: {e}")

    # ==========================================
    # WORD (DOCX) GENERATION
    # ==========================================
    def generate_docx_base64(self, kpi_metadata: Dict[str, Any], analysis_data: Dict[str, Any], language: str = "arabic") -> str:
        """Generate Word report and return as base64."""
        is_arabic = (language == "arabic")
        doc = Document()
        
        if is_arabic:
            self._set_rtl_document(doc)
            
        self._configure_docx_styles(doc)
        
        self._add_docx_title(doc, kpi_metadata, "تقرير تحليل المؤشر" if is_arabic else "KPI Analysis Report", is_arabic)
        doc.add_page_break()
        
        results = analysis_data.get(language, {})
        
        # Sections
        sections = [
            ("الملخص التنفيذي" if is_arabic else "Executive Summary", results.get("executive_summary", "")),
            ("تحليل الأداء" if is_arabic else "Performance Analysis", results.get("performance_analysis", "")),
            ("الأسباب الجذرية" if is_arabic else "Root Causes", results.get("root_causes", [])),
            ("التوصيات" if is_arabic else "Recommendations", results.get("recommendations", []))
        ]
        
        for title, content in sections:
            if content:
                self._add_docx_heading(doc, title, level=1, is_arabic=is_arabic)
                if isinstance(content, list):
                    for item in content:
                        self._add_docx_bullet(doc, item, is_arabic=is_arabic)
                else:
                    self._add_docx_text(doc, content, is_arabic=is_arabic)
                doc.add_paragraph()

        buffer = BytesIO()
        doc.save(buffer)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')

    def _set_rtl_document(self, doc):
        """Set document direction to RTL for Arabic."""
        for section in doc.sections:
            sectPr = section._sectPr
            if sectPr.find(qn('w:bidi')) is None:
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                sectPr.append(bidi)
        
        # Apply RTL to normal style
        if 'Normal' in doc.styles:
            style = doc.styles['Normal']
            pPr = style._element.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)

    def _configure_docx_styles(self, doc):
        """Configure standard styles."""
        font_name = 'Calibri'
        styles = doc.styles
        for style_name in ['Normal', 'Heading 1', 'Heading 2']:
            if style_name in styles:
                styles[style_name].font.name = font_name

    def _apply_arabic_run_properties(self, run):
        """Apply RTL and Arabic font run properties."""
        rPr = run._element.get_or_add_rPr()
        if rPr.find(qn('w:rtl')) is None:
            rtl = OxmlElement('w:rtl')
            rtl.set(qn('w:val'), '1')
            rPr.append(rtl)
        
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:cs'), 'Calibri')

    def _add_docx_title(self, doc, metadata, report_title, is_arabic):
        """Add title page content."""
        for _ in range(3): doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = self.docx_colors["primary"]
        if is_arabic: self._apply_arabic_run_properties(run)

        # Name
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kpi_name = metadata.get("kpi_name_ar") if is_arabic else (metadata.get("kpi_name_en") or metadata.get("kpi_name_ar"))
        run_name = p_name.add_run(kpi_name or "Unknown KPI")
        run_name.font.size = Pt(18)
        run_name.font.color.rgb = self.docx_colors["accent"]
        if is_arabic: self._apply_arabic_run_properties(run_name)
        
        # Entity
        entity = metadata.get("entity_name_ar") if is_arabic else (metadata.get("entity_name_en") or metadata.get("entity_name_ar", ""))
        if entity:
            p_ent = doc.add_paragraph()
            p_ent.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_ent = p_ent.add_run(entity)
            if is_arabic: self._apply_arabic_run_properties(run_ent)

    def _add_docx_heading(self, doc, text, level=1, is_arabic=False):
        heading = doc.add_heading(text, level=level)
        if is_arabic:
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pPr = heading._element.get_or_add_pPr()
            if pPr.find(qn('w:bidi')) is None:
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.insert(0, bidi)
            if heading.runs:
                self._apply_arabic_run_properties(heading.runs[0])
            heading.runs[0].font.color.rgb = self.docx_colors["primary"]
        return heading

    def _add_docx_text(self, doc, text, is_arabic=False):
        p = doc.add_paragraph()
        if is_arabic:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pPr = p._element.get_or_add_pPr()
            if pPr.find(qn('w:bidi')) is None:
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.insert(0, bidi)
        run = p.add_run(text)
        if is_arabic:
            self._apply_arabic_run_properties(run)
        return p

    def _add_docx_bullet(self, doc, text, is_arabic=False):
        p = doc.add_paragraph(text, style='List Bullet')
        if is_arabic:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pPr = p._element.get_or_add_pPr()
            if pPr.find(qn('w:bidi')) is None:
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.insert(0, bidi)
            if p.runs:
                self._apply_arabic_run_properties(p.runs[0])
        return p

    # ==========================================
    # PDF GENERATION
    # ==========================================
    def _reshape_arabic(self, text: str) -> str:
        """Reshape and reorder Arabic text for ReportLab rendering."""
        # Reshape to connect characters correctly
        reshaped = arabic_reshaper.reshape(text)
        # Apply bidi algorithm to reorder right-to-left
        bidi_text = get_display(reshaped)
        return bidi_text

    def generate_pdf_base64(self, kpi_metadata: Dict[str, Any], analysis_data: Dict[str, Any], language: str = "arabic") -> str:
        """Generate PDF report and return as base64."""
        is_arabic = (language == "arabic")
        buffer = BytesIO()
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=18
        )

        styles = getSampleStyleSheet()
        
        # Configure styles based on language
        from reportlab.lib.enums import TA_RIGHT, TA_LEFT, TA_CENTER
        title_align = TA_CENTER # Center
        body_align = TA_RIGHT if is_arabic else TA_LEFT # Right aligned for Arabic, Left for English
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=self.arabic_font_name,
            fontSize=24,
            textColor=HexColor('#374151'),
            alignment=title_align,
            spaceAfter=20
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=self.arabic_font_name,
            fontSize=16,
            textColor=HexColor('#4F46E5'),
            alignment=body_align,
            spaceAfter=12,
            spaceBefore=16,
            wordWrap='RTL' if is_arabic else 'LTR'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontName=self.arabic_font_name,
            fontSize=11,
            textColor=HexColor('#111827'),
            alignment=body_align,
            spaceAfter=8,
            leading=16,
            wordWrap='RTL' if is_arabic else 'LTR'
        )

        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=body_style,
            leftIndent=20 if not is_arabic else 0,
            rightIndent=20 if is_arabic else 0,
            bulletText='•',
            alignment=body_align
        )

        elements = []
        
        # Title Page
        report_title = "تقرير تحليل المؤشر" if is_arabic else "KPI Analysis Report"
        if is_arabic: report_title = self._reshape_arabic(report_title)
        elements.append(Paragraph(report_title, title_style))
        
        kpi_name = kpi_metadata.get("kpi_name_ar") if is_arabic else (kpi_metadata.get("kpi_name_en") or kpi_metadata.get("kpi_name_ar"))
        if is_arabic and kpi_name: kpi_name = self._reshape_arabic(kpi_name)
        elements.append(Paragraph(kpi_name or "Unknown KPI", heading_style))
        elements.append(HRFlowable(width="100%", thickness=1, color=HexColor('#E5E7EB'), spaceBefore=20, spaceAfter=20))
        
        # Sections
        results = analysis_data.get(language, {})
        sections = [
            ("الملخص التنفيذي" if is_arabic else "Executive Summary", results.get("executive_summary", "")),
            ("تحليل الأداء" if is_arabic else "Performance Analysis", results.get("performance_analysis", "")),
            ("الأسباب الجذرية" if is_arabic else "Root Causes", results.get("root_causes", [])),
            ("التوصيات" if is_arabic else "Recommendations", results.get("recommendations", []))
        ]
        
        for section_title, content in sections:
            if content:
                if is_arabic: section_title = self._reshape_arabic(section_title)
                elements.append(Paragraph(section_title, heading_style))
                
                if isinstance(content, list):
                    for item in content:
                        text = self._reshape_arabic(item) if is_arabic else item
                        # Bullet markup for ReportLab
                        elements.append(Paragraph(f"• {text}", bullet_style))
                else:
                    text = self._reshape_arabic(content) if is_arabic else content
                    elements.append(Paragraph(text, body_style))
                
                elements.append(Spacer(1, 12))

        doc.build(elements)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')

