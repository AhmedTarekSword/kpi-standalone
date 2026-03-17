import logging
import base64
import os
from io import BytesIO
from typing import Dict, Any, Optional
from datetime import datetime

# docx imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# reportlab imports for PDF
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.colors import HexColor
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_RIGHT, TA_LEFT, TA_CENTER
import arabic_reshaper
from bidi.algorithm import get_display

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Bilingual section labels
# ---------------------------------------------------------------------------
_SECTION_LABELS = {
    "arabic": {
        "report_title": "تقرير تحليل المؤشر",
        "executive_summary": "الملخص التنفيذي",
        "performance_analysis": "تحليل الأداء",
        "root_causes": "الأسباب الجذرية",
        "recommendations": "التوصيات",
    },
    "english": {
        "report_title": "KPI Analysis Report",
        "executive_summary": "Executive Summary",
        "performance_analysis": "Performance Analysis",
        "root_causes": "Root Causes",
        "recommendations": "Recommendations",
    },
}


class ReportGenerator:
    """Generate Word and PDF reports from KPI analysis with RTL/LTR support."""

    def __init__(self):
        self.colors_rgb = {
            "primary": (55, 65, 81),            # Charcoal gray
            "accent": (79, 70, 229),            # Modern indigo
            "success": (16, 185, 129),          # Fresh green
            "text_primary": (17, 24, 39),       # Near black
            "text_secondary": (107, 114, 128),  # Muted gray
        }

        self.docx_colors = {
            key: RGBColor(*rgb) for key, rgb in self.colors_rgb.items()
        }

        self._setup_pdf_fonts()

    def _setup_pdf_fonts(self):
        """Try to load a font that supports Arabic (like Arial or Tahoma) for PDF."""
        self.arabic_font_name = "Helvetica"  # Fallback

        font_paths = [
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/tahoma.ttf",
            "C:/Windows/Fonts/times.ttf",
        ]

        for path in font_paths:
            if os.path.exists(path):
                try:
                    font_name = os.path.basename(path).split(".")[0].capitalize()
                    pdfmetrics.registerFont(TTFont(font_name, path))
                    self.arabic_font_name = font_name
                    logger.info(f"Registered {font_name} for PDF generation.")
                    break
                except Exception as e:
                    logger.warning(f"Failed to load font {path}: {e}")

    # ==========================================
    # WORD (DOCX) GENERATION
    # ==========================================

    def generate_docx_base64(
        self,
        kpi_metadata: Dict[str, Any],
        analysis_data: Dict[str, Any],
        language: str = "arabic",
    ) -> str:
        """
        Generate Word report and return as base64.

        Args:
            language: "arabic", "english", or "both".
        """
        doc = Document()
        self._configure_docx_styles(doc)

        if language == "both":
            # Arabic section first (RTL)
            self._set_rtl_document(doc)
            self._add_docx_title(doc, kpi_metadata, _SECTION_LABELS["arabic"]["report_title"], is_arabic=True)
            doc.add_page_break()
            self._add_docx_language_sections(doc, analysis_data, "arabic")
            # English section after page break (LTR — achieved per-paragraph)
            doc.add_page_break()
            self._add_docx_title(doc, kpi_metadata, _SECTION_LABELS["english"]["report_title"], is_arabic=False)
            doc.add_page_break()
            self._add_docx_language_sections(doc, analysis_data, "english")
        else:
            is_arabic = language == "arabic"
            if is_arabic:
                self._set_rtl_document(doc)
            labels = _SECTION_LABELS[language]
            self._add_docx_title(doc, kpi_metadata, labels["report_title"], is_arabic=is_arabic)
            doc.add_page_break()
            self._add_docx_language_sections(doc, analysis_data, language)

        buffer = BytesIO()
        doc.save(buffer)
        return base64.b64encode(buffer.getvalue()).decode("utf-8")

    def _add_docx_language_sections(self, doc, analysis_data: Dict[str, Any], language: str):
        """Add the four report sections for a given language."""
        is_arabic = language == "arabic"
        labels = _SECTION_LABELS[language]
        results = analysis_data.get(language, {})

        sections = [
            (labels["executive_summary"], results.get("executive_summary", "")),
            (labels["performance_analysis"], results.get("performance_analysis", "")),
            (labels["root_causes"], results.get("root_causes", [])),
            (labels["recommendations"], results.get("recommendations", [])),
        ]

        for title, content in sections:
            if content:
                self._add_docx_heading(doc, title, level=1, is_arabic=is_arabic)
                if isinstance(content, list):
                    for item in content:
                        self._add_docx_bullet(doc, item, is_arabic=is_arabic)
                else:
                    self._add_docx_text(doc, content, is_arabic=is_arabic)
                doc.add_paragraph()

    def _set_rtl_document(self, doc):
        """Set document direction to RTL for Arabic."""
        for section in doc.sections:
            sectPr = section._sectPr
            if sectPr.find(qn("w:bidi")) is None:
                bidi = OxmlElement("w:bidi")
                bidi.set(qn("w:val"), "1")
                sectPr.append(bidi)

        if "Normal" in doc.styles:
            style = doc.styles["Normal"]
            pPr = style._element.get_or_add_pPr()
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            pPr.append(bidi)

    def _configure_docx_styles(self, doc):
        """Configure standard styles."""
        font_name = "Calibri"
        styles = doc.styles
        for style_name in ["Normal", "Heading 1", "Heading 2"]:
            if style_name in styles:
                styles[style_name].font.name = font_name

    def _apply_arabic_run_properties(self, run):
        """Apply RTL and Arabic font run properties."""
        rPr = run._element.get_or_add_rPr()
        if rPr.find(qn("w:rtl")) is None:
            rtl = OxmlElement("w:rtl")
            rtl.set(qn("w:val"), "1")
            rPr.append(rtl)

        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:cs"), "Calibri")

    def _add_docx_title(self, doc, metadata, report_title, is_arabic):
        """Add title page content."""
        for _ in range(3):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = self.docx_colors["primary"]
        if is_arabic:
            self._apply_arabic_run_properties(run)

        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kpi_name = (
            metadata.get("kpi_name_ar")
            if is_arabic
            else (metadata.get("kpi_name_en") or metadata.get("kpi_name_ar"))
        )
        run_name = p_name.add_run(kpi_name or "Unknown KPI")
        run_name.font.size = Pt(18)
        run_name.font.color.rgb = self.docx_colors["accent"]
        if is_arabic:
            self._apply_arabic_run_properties(run_name)

        entity = (
            metadata.get("entity_name_ar")
            if is_arabic
            else (metadata.get("entity_name_en") or metadata.get("entity_name_ar", ""))
        )
        if entity:
            p_ent = doc.add_paragraph()
            p_ent.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_ent = p_ent.add_run(entity)
            if is_arabic:
                self._apply_arabic_run_properties(run_ent)

    def _add_docx_heading(self, doc, text, level=1, is_arabic=False):
        heading = doc.add_heading(text, level=level)
        # if is_arabic:
        #     heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # else:
        #     heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = heading._element.get_or_add_pPr()
        if is_arabic and pPr.find(qn("w:bidi")) is None:
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            pPr.insert(0, bidi)
        if heading.runs:
            if is_arabic:
                self._apply_arabic_run_properties(heading.runs[0])
            heading.runs[0].font.color.rgb = self.docx_colors["primary"]
        return heading

    def _add_docx_text(self, doc, text, is_arabic=False):
        p = doc.add_paragraph()
        # p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_arabic else WD_ALIGN_PARAGRAPH.LEFT
        if is_arabic:
            pPr = p._element.get_or_add_pPr()
            if pPr.find(qn("w:bidi")) is None:
                bidi = OxmlElement("w:bidi")
                bidi.set(qn("w:val"), "1")
                pPr.insert(0, bidi)
        run = p.add_run(text)
        if is_arabic:
            self._apply_arabic_run_properties(run)
        return p

    def _add_docx_bullet(self, doc, text, is_arabic=False):
        p = doc.add_paragraph(text, style="List Bullet")
        # p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_arabic else WD_ALIGN_PARAGRAPH.LEFT
        if is_arabic:
            pPr = p._element.get_or_add_pPr()
            if pPr.find(qn("w:bidi")) is None:
                bidi = OxmlElement("w:bidi")
                bidi.set(qn("w:val"), "1")
                pPr.insert(0, bidi)
            if p.runs:
                self._apply_arabic_run_properties(p.runs[0])
        return p

    # ==========================================
    # PDF GENERATION
    # ==========================================

    def _reshape_arabic(self, text: str) -> str:
        """Reshape and reorder Arabic text for ReportLab rendering."""
        reshaped = arabic_reshaper.reshape(text)
        return get_display(reshaped)

    def _build_pdf_styles(self, is_arabic: bool):
        """Return a dict of ReportLab ParagraphStyle objects for the given direction."""
        styles = getSampleStyleSheet()
        body_align = TA_RIGHT if is_arabic else TA_LEFT

        title_style = ParagraphStyle(
            f"CustomTitle_{'ar' if is_arabic else 'en'}",
            parent=styles["Heading1"],
            fontName=self.arabic_font_name,
            fontSize=24,
            textColor=HexColor("#374151"),
            alignment=TA_CENTER,
            spaceAfter=20,
        )

        heading_style = ParagraphStyle(
            f"CustomHeading_{'ar' if is_arabic else 'en'}",
            parent=styles["Heading2"],
            fontName=self.arabic_font_name,
            fontSize=16,
            textColor=HexColor("#4F46E5"),
            alignment=body_align,
            spaceAfter=12,
            spaceBefore=16,
            wordWrap="RTL" if is_arabic else "LTR",
        )

        body_style = ParagraphStyle(
            f"CustomBody_{'ar' if is_arabic else 'en'}",
            parent=styles["Normal"],
            fontName=self.arabic_font_name,
            fontSize=11,
            textColor=HexColor("#111827"),
            alignment=body_align,
            spaceAfter=8,
            leading=16,
            wordWrap="RTL" if is_arabic else "LTR",
        )

        bullet_style = ParagraphStyle(
            f"CustomBullet_{'ar' if is_arabic else 'en'}",
            parent=body_style,
            leftIndent=20 if not is_arabic else 0,
            rightIndent=20 if is_arabic else 0,
            bulletText="•",
            alignment=body_align,
        )

        return {
            "title": title_style,
            "heading": heading_style,
            "body": body_style,
            "bullet": bullet_style,
        }

    def _add_pdf_language_sections(
        self,
        elements: list,
        analysis_data: Dict[str, Any],
        language: str,
    ):
        """Append PDF story elements for a single language."""
        is_arabic = language == "arabic"
        pdf_styles = self._build_pdf_styles(is_arabic)
        labels = _SECTION_LABELS[language]
        results = analysis_data.get(language, {})

        sections = [
            (labels["executive_summary"], results.get("executive_summary", "")),
            (labels["performance_analysis"], results.get("performance_analysis", "")),
            (labels["root_causes"], results.get("root_causes", [])),
            (labels["recommendations"], results.get("recommendations", [])),
        ]

        for section_title, content in sections:
            if content:
                heading_text = self._reshape_arabic(section_title) if is_arabic else section_title
                elements.append(Paragraph(heading_text, pdf_styles["heading"]))

                if isinstance(content, list):
                    for item in content:
                        text = self._reshape_arabic(item) if is_arabic else item
                        elements.append(Paragraph(f"• {text}", pdf_styles["bullet"]))
                else:
                    text = self._reshape_arabic(content) if is_arabic else content
                    elements.append(Paragraph(text, pdf_styles["body"]))

                elements.append(Spacer(1, 12))

    def generate_pdf_base64(
        self,
        kpi_metadata: Dict[str, Any],
        analysis_data: Dict[str, Any],
        language: str = "arabic",
    ) -> str:
        """
        Generate PDF report and return as base64.

        Args:
            language: "arabic", "english", or "both".
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        elements = []
        languages_to_render = ["arabic", "english"] if language == "both" else [language]

        for idx, lang in enumerate(languages_to_render):
            is_arabic = lang == "arabic"
            pdf_styles = self._build_pdf_styles(is_arabic)
            labels = _SECTION_LABELS[lang]

            # Title
            report_title = labels["report_title"]
            if is_arabic:
                report_title = self._reshape_arabic(report_title)
            elements.append(Paragraph(report_title, pdf_styles["title"]))

            kpi_name = (
                kpi_metadata.get("kpi_name_ar")
                if is_arabic
                else (kpi_metadata.get("kpi_name_en") or kpi_metadata.get("kpi_name_ar"))
            )
            if is_arabic and kpi_name:
                kpi_name = self._reshape_arabic(kpi_name)
            elements.append(Paragraph(kpi_name or "Unknown KPI", pdf_styles["heading"]))
            elements.append(
                HRFlowable(
                    width="100%",
                    thickness=1,
                    color=HexColor("#E5E7EB"),
                    spaceBefore=20,
                    spaceAfter=20,
                )
            )

            self._add_pdf_language_sections(elements, analysis_data, lang)

            # Separator between language sections in "both" mode
            if language == "both" and idx == 0:
                elements.append(Spacer(1, 36))
                elements.append(
                    HRFlowable(
                        width="100%",
                        thickness=2,
                        color=HexColor("#4F46E5"),
                        spaceBefore=10,
                        spaceAfter=30,
                    )
                )

        doc.build(elements)
        return base64.b64encode(buffer.getvalue()).decode("utf-8")
